#!/usr/bin/env python3
"""Generate the BigQuery Conversational Agent API Technical Design Document (DOCX)."""

from datetime import date
from pathlib import Path
import sys

SCRIPTS_DIR = Path(__file__).resolve().parent
sys.path.insert(0, str(SCRIPTS_DIR))

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from generate_diagrams import generate_all


OUTPUT = Path(__file__).resolve().parent.parent / "docs" / "BQ-Agent-Technical-Design.docx"


def set_cell_shading(cell, fill: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shading)


def add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_para(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_code(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], header_fill: str = "D9E2F3") -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_shading(hdr[i], header_fill)
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            table.rows[ri + 1].cells[ci].text = val
    doc.add_paragraph()


def add_image(doc: Document, path: Path, width_inches: float = 6.5, caption: str | None = None) -> None:
    if not path.exists():
        add_para(doc, f"[Missing diagram: {path}]", bold=True)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width_inches))
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = cap.add_run(caption)
        cr.italic = True
        cr.font.size = Pt(9)
        cr.font.color.rgb = RGBColor(0x5F, 0x63, 0x68)
    doc.add_paragraph()


def build_document(diagram_paths: dict[str, Path]) -> Document:
    doc = Document()
    today = date.today().isoformat()

    add_title(doc, "BigQuery Conversational Agent API")
    add_title(doc, "Technical Design Document")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Version 1.0  |  {today}")
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x5F, 0x63, 0x68)
    doc.add_page_break()

    # 1. Executive Summary
    add_heading(doc, "1. Executive Summary")
    add_para(
        doc,
        "This document describes the technical architecture of the BigQuery Conversational Agent "
        "Chatbot API: an application that exposes REST APIs for multi-turn conversational "
        "analytics over a pre-created BigQuery Gemini Data Agent (GDA). The API combines three Google "
        "Cloud capabilities—Vertex AI Sessions (short-term conversation history), Vertex AI Memory Bank "
        "(long-term user context), and the Conversational Analytics API stateless :chat endpoint—to "
        "deliver stateful, personalized data-agent chat without relying on GDA conversation references. "
        "User feedback is persisted to BigQuery for observability and iterative agent improvement."
    )

    # 2. Goals
    add_heading(doc, "2. Goals and Non-Goals")
    add_heading(doc, "2.1 Goals", level=2)
    add_bullets(doc, [
        "Provide a REST API for stateful multi-turn chat with a BigQuery data agent.",
        "Persist conversation turns in Vertex AI Sessions scoped by conversationId.",
        "Retrieve and generate long-term user memories via Vertex AI Memory Bank scoped by userId.",
        "Support streaming (SSE) and non-streaming chat responses.",
        "Collect end-user feedback (thumbs up/down, comments) in BigQuery.",
        "Authenticate to GCP using Application Default Credentials (service account or user ADC).",
    ])
    add_heading(doc, "2.2 Non-Goals", level=2)
    add_bullets(doc, [
        "Creating or editing BigQuery data agents at runtime (agents are pre-created in BigQuery).",
        "Automatic agent retraining from feedback (feedback is stored for human review, similar to Snowflake Cortex).",
        "Managing GDA conversation references (each chat turn uses stateless :chat with replayed history).",
    ])

    # 3. Architecture
    add_heading(doc, "3. System Architecture")
    add_heading(doc, "3.1 High-Level Architecture", level=2)
    add_para(
        doc,
        "Figure 1 shows how the system is split into a BigQuery data plane (knowledge, queries, "
        "feedback) and a live chat path that orchestrates Vertex AI context with stateless Gemini "
        "Data Agent (GDA) calls. The GDA itself is hosted in BigQuery—it is created and managed "
        "as a BigQuery resource, not as a separate Vertex AI agent. The application sits in the "
        "middle: it loads history and memories, replays them on each turn, and persists results "
        "back to GCP."
    )
    add_image(
        doc,
        diagram_paths["architecture"],
        width_inches=6.8,
        caption="Figure 1 — High-level system architecture",
    )

    add_heading(doc, "3.1.1 BigQuery Data Plane (left panel)", level=3)
    add_para(
        doc,
        "The left panel is provisioned ahead of time and does not participate in every HTTP request "
        "directly. It holds the assets the data agent needs to answer questions—including the "
        "GDA resource hosted in BigQuery—and the table where user feedback is stored."
    )
    add_bullets(doc, [
        "Knowledge & agent — Source tables, verified SQL queries, and a pre-created Gemini Data "
        "Agent (GDA) hosted in BigQuery and registered against that knowledge. The agent is "
        "created and configured in BigQuery (e.g. via bq mk --data_agent); the application only "
        "references it by dataAgentId when calling the Conversational Analytics API.",
        "BigQuery (SQL) — The analytical store the GDA queries at runtime when translating natural "
        "language into SQL and executing it against live data.",
        "Feedback store (agent_feedback.agent_feedback) — A BigQuery table that receives thumbs "
        "up/down, comments, and metadata from POST /conversations/{id}/feedback for observability "
        "and offline review.",
    ])

    add_heading(doc, "3.1.2 Live Chat Path (right panel)", level=3)
    add_para(
        doc,
        "The right panel is the per-request path for conversational analytics. Components flow "
        "left to right; dashed arrows show secondary or return flows."
    )
    add_table(
        doc,
        ["Component", "Role"],
        [
            ["Client", "Calls REST endpoints (including SSE streaming) to create conversations and send messages."],
            ["Application API", "Orchestrates each turn: loads context, builds the :chat payload, streams or returns the answer, persists the turn, and triggers async memory generation."],
            ["Vertex AI (Sessions + Memory Bank)", "Sessions store short-term turn history keyed by conversationId. Memory Bank stores long-term user facts keyed by userId, retrieved before chat and updated after."],
            ["Conversational Analytics API", "Google Cloud service that exposes the stateless DataChatService.chat endpoint used for every agent turn."],
            ["Data Agent (NL → SQL)", "The pre-created GDA hosted in BigQuery. It converts the user question (plus replayed context) into SQL and executes it against BigQuery tables. Invoked via the Conversational Analytics API using the agent's dataAgentId."],
            ["Live data", "BigQuery tables queried by the agent during the turn; results flow back through the pipeline to the client."],
        ],
    )

    add_heading(doc, "3.1.3 Flows Shown in the Diagram", level=3)
    add_bullets(doc, [
        "Primary flow (solid arrows, left to right) — Client → Application → Vertex AI → Conversational Analytics API → Data Agent → Live data. This is the main path for each chat message.",
        "Load context (dashed, upward) — Before calling :chat, the application reads session history and Memory Bank entries for the conversation/user and includes them in the request.",
        "Query (solid, downward) — The data agent executes generated SQL against BigQuery live data.",
        "Answer (dashed, return) — The agent response is streamed or returned through the application to the client.",
        "Feedback (dashed, to data plane) — Optional user feedback is written asynchronously to the agent_feedback table in BigQuery.",
    ])
    add_para(
        doc,
        "Each turn follows the same pattern: load session history and memories → call stateless :chat "
        "with full context replayed → persist the user and agent messages to Sessions → asynchronously "
        "invoke memories:generate so Memory Bank can extract durable user context for future turns. "
        "See Section 3.2 for the step-by-step sequence."
    )

    add_heading(doc, "3.2 Conversation Chat Sequence", level=2)
    add_para(
        doc,
        "Figure 2 walks through a single stateful turn: POST /conversations/{conversationId}/messages "
        "(or the SSE streaming variant). The diagram is divided into three phases—load context, call the "
        "BigQuery-hosted GDA, then persist the turn—reflecting the order of operations in the application."
    )
    add_image(
        doc,
        diagram_paths["sequence"],
        width_inches=6.8,
        caption="Figure 2 — Conversation chat sequence",
    )

    add_heading(doc, "3.2.1 Participants", level=3)
    add_para(
        doc,
        "Five actors appear in the sequence. Vertex AI components manage conversation state and long-term "
        "memory; the GDA is invoked through the Conversational Analytics API but remains a BigQuery resource."
    )
    add_table(
        doc,
        ["Actor", "Role in this request"],
        [
            ["Client", "Sends the user message via REST (JSON response or SSE stream)."],
            ["Application API", "Orchestrates all phases: reads context, builds the ChatRequest, streams the GDA response, then writes to Sessions and Memory Bank."],
            ["Vertex Sessions", "Stores prior turns for the conversationId. Events are read before chat and appended after the agent completes."],
            ["Memory Bank", "Returns relevant long-term user facts (memories:retrieve) before chat; asynchronously extracts new facts (memories:generate) after the turn is persisted."],
            ["GDA :chat", "The BigQuery-hosted Gemini Data Agent, called statelessly via DataChatService.chat. Each request includes replayed history, memories, and the new user message."],
        ],
    )

    add_heading(doc, "3.2.2 Phase 1 — Load context", level=3)
    add_para(
        doc,
        "Before any agent call, the application gathers everything needed to make the turn stateful and "
        "personalized. No GDA or BigQuery query runs in this phase."
    )
    add_bullets(doc, [
        "POST /messages — Client submits the new user message (and optionally userId) for an existing conversation.",
        "GET session events — Application lists all prior Session events for conversationId and maps them to GDA message history.",
        "Conversation history — Prior user and agent turns are returned from Vertex Sessions and will be replayed in the ChatRequest.",
        "memories:retrieve — Application queries Memory Bank for the userId (top-K, similarity-based) using the current message as the retrieval query.",
        "User facts — Retrieved memories are formatted into a context block prepended to the user message so the agent can personalize its answer.",
    ])

    add_heading(doc, "3.2.3 Phase 2 — Agent call", level=3)
    add_para(
        doc,
        "The application builds a stateless ChatRequest and streams the response back to the client as "
        "it arrives. The GDA runs in BigQuery; the Conversational Analytics API is the transport layer."
    )
    add_bullets(doc, [
        "ChatRequest (history + memories + message) — Application sends replayed session history, the memory-augmented user message, and the dataAgentId of the BigQuery-hosted GDA to DataChatService.chat.",
        "Response stream chunks — GDA translates natural language to SQL, executes against BigQuery source tables, and streams structured response messages back to the application.",
        "SSE / JSON chunks to client — Application forwards chunks to the client immediately (SSE for /messages/stream, aggregated JSON for /messages) so the user sees progress before persistence completes.",
    ])

    add_heading(doc, "3.2.4 Phase 3 — Persist turn (after agent completes)", level=3)
    add_para(
        doc,
        "Session and Memory Bank updates happen only after the full GDA response has been received. "
        "This avoids recording partial or failed agent turns in conversation history."
    )
    add_bullets(doc, [
        "appendEvent — user message — Application writes the original user text to Vertex Sessions with a new invocationId linking the turn.",
        "appendEvent — agent response — Application writes the final agent text (extracted from streamed chunks) to Sessions under the same invocationId.",
        "memories:generate (async) — Application triggers Memory Bank asynchronously with the updated session events so durable user facts can be extracted for future turns. Failures here do not affect the chat response already returned to the client.",
    ])
    add_para(
        doc,
        "Important design choices reflected in this sequence: (1) GDA :chat is stateless—conversation "
        "continuity comes from replaying Vertex Session history on every turn, not from a GDA conversation "
        "reference; (2) the GDA is hosted in BigQuery and identified by dataAgentId; (3) the client "
        "receives the streamed answer before Sessions are updated, but the next turn will include this "
        "turn once appendEvent completes."
    )

    add_heading(doc, "3.3 Application Integration", level=2)
    add_table(
        doc,
        ["Layer", "Technology"],
        [
            ["Application", "REST API layer (conversations, chat, feedback, memories)"],
            ["GCP clients", "Gemini Data Analytics, Vertex AI Platform, BigQuery APIs"],
            ["GDA integration", "Conversational Analytics API — DataChatService.chat"],
            ["Vertex integration", "Sessions and Memory Bank REST APIs"],
            ["BigQuery integration", "Streaming insert for feedback; agent queries source data"],
            ["Auth", "Google Application Default Credentials (cloud-platform scope)"],
        ],
    )

    # 4. Capabilities
    add_heading(doc, "4. Capability Breakdown")
    add_para(doc, "Figure 3 maps each API capability to its GCP service and scope key:")
    add_image(
        doc,
        diagram_paths["capabilities"],
        width_inches=6.5,
        caption="Figure 3 — Capability to GCP service mapping",
    )
    add_table(
        doc,
        ["Capability", "GCP Service", "API / Endpoint", "Scope Key"],
        [
            ["Session management", "Vertex AI Agent Engine (Sessions)", "aiplatform.googleapis.com v1beta1", "conversationId"],
            ["Memory Bank", "Vertex AI Agent Engine (Memory Bank)", "memories:retrieve, memories:generate", "userId"],
            ["Stateless GDA chat", "Gemini Data Analytics", "geminidataanalytics.googleapis.com DataChatService.chat", "dataAgentId"],
            ["One-off chat", "Gemini Data Analytics", "Same :chat, no session/memory", "N/A"],
            ["User feedback", "BigQuery", "insertAll streaming insert", "conversationId + invocationId"],
            ["List memories", "Vertex AI Memory Bank", "GET memories (list) + filter", "userId"],
        ],
    )

    # 5. APIs to enable
    add_heading(doc, "5. Google Cloud APIs to Enable")
    add_para(doc, "Enable the following APIs in the target GCP project before deployment:")
    add_code(
        doc,
        """gcloud config set project YOUR_PROJECT_ID

# Core data agent chat
gcloud services enable geminidataanalytics.googleapis.com
gcloud services enable cloudaicompanion.googleapis.com
gcloud services enable bigquery.googleapis.com

# Vertex AI Sessions & Memory Bank
gcloud services enable aiplatform.googleapis.com

# Optional: Knowledge Catalog (BigQuery Agent Catalog metadata)
gcloud services enable cloudasset.googleapis.com"""
    )
    add_table(
        doc,
        ["API", "Service name", "Used by"],
        [
            ["Gemini Data Analytics API", "geminidataanalytics.googleapis.com", "Stateless :chat with data agent"],
            ["Cloud AI Companion API", "cloudaicompanion.googleapis.com", "Gemini for Google Cloud companion features"],
            ["BigQuery API", "bigquery.googleapis.com", "Agent queries data sources; feedback storage"],
            ["Vertex AI API", "aiplatform.googleapis.com", "Sessions, Memory Bank, Reasoning Engines"],
        ],
    )

    # 6. IAM
    add_heading(doc, "6. IAM Roles and Permissions by Capability")
    add_para(
        doc,
        "Grant roles to the runtime service account (e.g. sa-bq-agent@PROJECT.iam.gserviceaccount.com) "
        "referenced by GOOGLE_APPLICATION_CREDENTIALS."
    )

    add_heading(doc, "6.1 Session Management", level=2)
    add_table(
        doc,
        ["Permission / Role", "Purpose"],
        [
            ["roles/aiplatform.user", "Create/get/delete sessions; list/append session events"],
            ["Reasoning Engine resource", "Empty Agent Engine instance (sessions-engine-id)"],
        ],
    )
    add_para(doc, "Vertex REST operations used:")
    add_bullets(doc, [
        "POST .../reasoningEngines/{id}/sessions?sessionId={conversationId}",
        "GET .../sessions/{conversationId}",
        "GET .../sessions/{conversationId}/events",
        "POST .../sessions/{conversationId}:appendEvent",
        "DELETE .../sessions/{conversationId}",
    ])

    add_heading(doc, "6.2 Memory Bank", level=2)
    add_table(
        doc,
        ["Permission / Role", "Purpose"],
        [
            ["roles/aiplatform.user", "Retrieve and generate memories"],
            ["Reasoning Engine resource", "Separate Memory Bank engine (memory-bank-engine-id)"],
        ],
    )
    add_bullets(doc, [
        "POST .../reasoningEngines/{id}/memories:retrieve — before chat (topK=5, similarity or simple)",
        "POST .../reasoningEngines/{id}/memories:generate — after chat (async, from session events)",
        "GET .../reasoningEngines/{id}/memories?filter=scope.user_id=\"{userId}\" — list all memories",
    ])

    add_heading(doc, "6.3 Stateless Conversational Analytics (GDA :chat)", level=2)
    add_table(
        doc,
        ["Role", "Purpose"],
        [
            ["roles/geminidataanalytics.dataAgentUser", "Invoke data agent in conversational mode"],
            ["roles/geminidataanalytics.dataAgentStatelessUser", "Required for stateless :chat (history replayed by API)"],
            ["roles/cloudaicompanion.user", "Cloud AI Companion access for Gemini features"],
            ["roles/bigquery.jobUser", "Run queries on behalf of the data agent"],
            ["BigQuery dataset: roles/bigquery.dataViewer", "Read knowledge-source tables the agent queries"],
        ],
    )
    add_para(doc, "GDA parent resource: projects/{project}/locations/{gcp-location} (default: us)")
    add_para(doc, "Data agent resource: projects/{project}/locations/{location}/dataAgents/{data-agent-id}")

    add_heading(doc, "6.4 Agent Feedback (BigQuery)", level=2)
    add_table(
        doc,
        ["Role", "Purpose"],
        [
            ["Dataset WRITER or roles/bigquery.dataEditor on agent_feedback", "Streaming insert feedback rows"],
            ["roles/bigquery.admin (optional)", "Auto-create dataset/table when feedback-auto-create-table=true"],
        ],
    )

    add_heading(doc, "6.5 Data Agent Creation (Operators / Data Engineers)", level=2)
    add_para(doc, "These roles are required for humans creating the agent in BigQuery, not the runtime SA:")
    add_table(
        doc,
        ["Role", "Purpose"],
        [
            ["roles/geminidataanalytics.dataAgentCreator", "Create data agents in BigQuery Agent Catalog"],
            ["roles/geminidataanalytics.dataAgentOwner", "Edit, publish, share, delete agents"],
            ["roles/geminidataanalytics.dataAgentEditor", "Edit agent configuration and verified queries"],
            ["roles/bigquery.admin or dataOwner", "Create source datasets/tables"],
        ],
    )

    add_heading(doc, "6.6 IAM Summary Matrix", level=2)
    add_table(
        doc,
        ["Capability", "Runtime SA roles", "Additional scope"],
        [
            ["Sessions", "aiplatform.user", "Sessions reasoning engine ID"],
            ["Memory Bank", "aiplatform.user", "Memory Bank reasoning engine ID"],
            ["GDA stateless chat", "dataAgentUser, dataAgentStatelessUser, cloudaicompanion.user, bigquery.jobUser", "dataViewer on source datasets"],
            ["Feedback", "bigquery.dataEditor", "agent_feedback dataset"],
            ["Create agent (ops)", "dataAgentCreator/Owner", "BigQuery console or API"],
        ],
    )

    # 7. REST API
    add_heading(doc, "7. REST API Specification")
    add_table(
        doc,
        ["Method", "Path", "Description"],
        [
            ["GET", "/health", "Health check"],
            ["POST", "/api/v1/conversations", "Create Vertex Session (optional conversationId, userId)"],
            ["DELETE", "/api/v1/conversations/{id}", "Delete session"],
            ["GET", "/api/v1/conversations/{id}/messages", "List session events"],
            ["POST", "/api/v1/conversations/{id}/messages", "Chat — JSON response"],
            ["POST", "/api/v1/conversations/{id}/messages/stream", "Chat — SSE stream, ends with [DONE]"],
            ["GET", "/api/v1/users/{userId}/memories", "List Memory Bank facts"],
            ["POST", "/api/v1/conversations/{id}/feedback", "Submit feedback → BigQuery"],
            ["POST", "/api/v1/chat", "One-off stateless chat (no session/memory)"],
            ["POST", "/api/v1/chat/stream", "One-off stateless chat (SSE)"],
        ],
    )

    add_heading(doc, "7.1 Feedback Request Body", level=2)
    add_table(
        doc,
        ["Field", "Required", "Description"],
        [
            ["positive", "Yes", "true = thumbs up, false = thumbs down"],
            ["feedbackMessage", "No", "Free-text comment (max 4096 chars)"],
            ["categories", "No", "Array of category strings"],
            ["invocationId", "No", "Turn ID from session messages"],
            ["userId", "No", "Overrides X-User-Id header"],
        ],
    )

    # 8. Configuration
    add_heading(doc, "8. Configuration Properties")
    add_table(
        doc,
        ["Property", "Env var", "Default", "Description"],
        [
            ["app.gcp-project-id", "GCP_PROJECT_ID", "bq-agent-poc-500313", "GCP project"],
            ["app.gcp-location", "GCP_LOCATION", "us", "GDA API location"],
            ["app.data-agent-id", "DATA_AGENT_ID", "(see application.yml)", "Pre-created BigQuery data agent ID"],
            ["app.vertex-location", "VERTEX_LOCATION", "us-central1", "Vertex Sessions/Memory Bank region"],
            ["app.sessions-engine-id", "VERTEX_SESSIONS_ENGINE_ID", "(required)", "Sessions reasoning engine"],
            ["app.memory-bank-engine-id", "VERTEX_MEMORY_BANK_ENGINE_ID", "(required)", "Memory Bank engine"],
            ["app.default-user-id", "DEFAULT_USER_ID", "default-user", "Fallback userId"],
            ["app.feedback-dataset", "FEEDBACK_DATASET", "agent_feedback", "Feedback BQ dataset"],
            ["app.feedback-table", "FEEDBACK_TABLE", "agent_feedback", "Feedback BQ table"],
            ["app.feedback-auto-create-table", "FEEDBACK_AUTO_CREATE_TABLE", "true", "Auto-create BQ resources"],
        ],
    )

    # 9. BigQuery Data Agent Creation
    add_heading(doc, "9. BigQuery Data Agent Creation with Verified Queries")
    add_para(
        doc,
        "The application does not create data agents at runtime. Agents are created in the BigQuery Agent Catalog "
        "(Google Cloud console) or via the Gemini Data Analytics API, then referenced by DATA_AGENT_ID. "
        "Verified queries (formerly golden queries) are SQL templates the agent invokes exactly when a user "
        "question matches, ensuring verifiable, non-hallucinated answers grounded in BigQuery data."
    )

    add_heading(doc, "9.1 Prerequisites", level=2)
    add_bullets(doc, [
        "Enable APIs listed in Section 5.",
        "Grant dataAgentCreator or dataAgentOwner to the operator.",
        "Prepare knowledge-source tables with business data.",
        "Run and validate SQL in BigQuery Studio before adding as verified queries.",
    ])

    add_heading(doc, "9.2 Step 1 — Create Source Dataset and Table", level=2)
    add_para(doc, "Example: used vehicle inventory (matches POC demo agent). Run in BigQuery:")
    add_code(
        doc,
        """-- scripts/create_data_agent_source.sql (excerpt)
CREATE SCHEMA IF NOT EXISTS `inventory`
OPTIONS (description = 'Used vehicle inventory for conversational analytics');

CREATE TABLE IF NOT EXISTS `inventory.used_cars` (
  brand STRING,
  model STRING,
  model_year INT64,
  milage STRING,
  fuel_type STRING,
  transmission STRING,
  ext_col STRING,
  int_col STRING,
  accident STRING,
  clean_title BOOL,
  price INT64
);

-- Load your CSV or INSERT sample rows, then verify:
SELECT brand, COUNT(*) AS cnt, AVG(price) AS avg_price
FROM `inventory.used_cars`
GROUP BY brand
ORDER BY cnt DESC;"""
    )

    add_heading(doc, "9.3 Step 2 — Create Agent in BigQuery Agent Catalog", level=2)
    add_bullets(doc, [
        "Navigate to BigQuery → Agents → Create custom agent.",
        "Name: e.g. Used Car Inventory Agent.",
        "Add knowledge source: inventory.used_cars.",
        "Add system instructions describing the assistant role and domain vocabulary.",
        "Publish the agent.",
        "Copy the agent ID (e.g. agent_c2454530-1ce3-47ee-a335-31ef826b3227) into DATA_AGENT_ID.",
    ])

    add_heading(doc, "9.4 Step 3 — Add Verified Queries", level=2)
    add_para(
        doc,
        "For each verified query: (1) enter the natural-language question, (2) write or Generate SQL, "
        "(3) Run and verify results in BigQuery Studio, (4) Add to the agent. "
        "When matched, the agent runs the SQL exactly as written."
    )

    add_heading(doc, "Verified Query 1 — Toyota cars under a price threshold", level=3)
    add_para(doc, "Question: \"Show me Toyota cars under $20k\"")
    add_code(
        doc,
        """SELECT
  brand, model, model_year, milage, fuel_type, transmission,
  ext_col, int_col, accident, clean_title, price
FROM `inventory.used_cars`
WHERE LOWER(brand) = 'toyota'
  AND price < 20000
ORDER BY price ASC;"""
    )

    add_heading(doc, "Verified Query 2 — Average price by brand", level=3)
    add_para(doc, "Question: \"What is the average price by brand?\"")
    add_code(
        doc,
        """SELECT
  brand,
  COUNT(*) AS vehicle_count,
  ROUND(AVG(price), 2) AS avg_price,
  MIN(price) AS min_price,
  MAX(price) AS max_price
FROM `inventory.used_cars`
GROUP BY brand
ORDER BY avg_price DESC;"""
    )

    add_heading(doc, "Verified Query 3 — Hybrids with clean title", level=3)
    add_para(doc, "Question: \"Are there any Toyota hybrids with a clean title?\"")
    add_code(
        doc,
        """SELECT
  brand, model, model_year, fuel_type, price, accident, clean_title
FROM `inventory.used_cars`
WHERE LOWER(brand) = 'toyota'
  AND LOWER(fuel_type) LIKE '%hybrid%'
  AND clean_title = TRUE
ORDER BY price ASC;"""
    )

    add_heading(doc, "Verified Query 4 — Low mileage vehicles (parameterized pattern)", level=3)
    add_para(doc, "Question template: \"Show me {brand} cars with less than {miles} miles\"")
    add_code(
        doc,
        """-- Parameterized verified query example (configure @brand, @max_miles in Agent Catalog)
SELECT brand, model, model_year, milage, price
FROM `inventory.used_cars`
WHERE LOWER(brand) = LOWER(@brand)
  AND SAFE_CAST(REGEXP_REPLACE(milage, r'[^0-9]', '') AS INT64) < @max_miles
ORDER BY price ASC;"""
    )

    add_heading(doc, "9.5 Step 4 — Validate Agent Before API Integration", level=2)
    add_bullets(doc, [
        "Chat with the agent in BigQuery Studio; confirm SQL is shown and results match manual queries.",
        "Test multi-turn questions that rely on verified query patterns.",
        "Note the agent ID and set DATA_AGENT_ID in application.yml.",
        "Test via API: POST /api/v1/chat with a simple question, then POST /api/v1/conversations for stateful flow.",
    ])

    # 10. Vertex setup
    add_heading(doc, "10. Vertex AI Agent Engine Setup")
    add_para(doc, "Run scripts/create_agent_engines.py to create empty Reasoning Engine instances:")
    add_code(
        doc,
        """source .venv/bin/activate
pip install "google-cloud-aiplatform>=1.111.0"
export GOOGLE_APPLICATION_CREDENTIALS=/path/to/sa-key.json
python scripts/create_agent_engines.py

# Output:
# SESSIONS_ENGINE_ID=5090605245940105216
# MEMORY_BANK_ENGINE_ID=4174122721770209280"""
    )

    # 11. Feedback table
    add_heading(doc, "11. Feedback Table Setup")
    add_code(
        doc,
        """bq query --project_id="$GCP_PROJECT_ID" --use_legacy_sql=false \\
  < scripts/create_agent_feedback_table.sql

-- Grant runtime SA write access:
-- bq update --source access.json PROJECT:agent_feedback
-- access.json: {"access":[{"role":"WRITER","userByEmail":"sa-bq-agent@PROJECT.iam.gserviceaccount.com"}]}"""
    )

    # 12. Security
    add_heading(doc, "12. Security Considerations")
    add_bullets(doc, [
        "Service account keys should be stored securely; prefer Workload Identity on Cloud Run/GKE.",
        "Scope BigQuery dataViewer to only datasets the agent needs.",
        "Scope bigquery.dataEditor to only the agent_feedback dataset.",
        "userId and conversationId are client-supplied; validate in production gateways.",
        "Feedback text may contain PII — apply BigQuery column-level security or DLP as needed.",
    ])

    # 13. References
    add_heading(doc, "13. References")
    add_bullets(doc, [
        "Conversational Analytics API: https://cloud.google.com/gemini/docs/conversational-analytics-api/overview",
        "Create data agents: https://cloud.google.com/bigquery/docs/create-data-agents",
        "Enable Conversational Analytics API: https://cloud.google.com/gemini/data-agents/conversational-analytics-api/enable-the-api",
        "Agent Platform Sessions: https://cloud.google.com/gemini/docs/agent-platform/sessions/manage-with-adk",
        "Memory Bank setup: https://cloud.google.com/gemini/docs/agent-platform/memory-bank/setup",
    ])

    return doc


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    diagram_paths = generate_all()
    doc = build_document(diagram_paths)
    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")
    for name, path in diagram_paths.items():
        print(f"  diagram/{name}: {path}")


if __name__ == "__main__":
    main()
